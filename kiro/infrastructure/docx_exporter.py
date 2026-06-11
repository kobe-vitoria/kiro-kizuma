"""Exporta ArticleDraft para .docx (Microsoft Word / Google Docs compatível).

Use case: enquanto o time não tem permissão no Confluence, os drafts saem como
.docx pra serem subidos no Google Drive e revisados pela equipe via comentários.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from kiro.domain.models import ArticleDraft, Cluster
from kiro.utils.branding import SIGNATURE

_LEADING_NUMBER_RE = re.compile(r"^\s*\d+\s*[.\)\-]\s+")


def article_to_docx(article: ArticleDraft, cluster: Cluster, output_path: Path) -> Path:
    """Renderiza um draft como arquivo .docx. Retorna o path escrito.

    Formatação:
      - Título grande em destaque
      - Subtítulo itálico com contagem de tickets
      - Tickets de origem em parágrafo bold-prefix
      - Seções H1 (Problema, Causa raiz, Solução, FAQ, Metadados)
      - Solução como lista numerada nativa do Word
      - FAQ como pares pergunta(bold)/resposta
      - Rodapé com slogan KIRO centralizado em cinza
    """
    doc = Document()

    # Fonte padrão limpa
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ─── Cabeçalho ────────────────────────────────────────────────
    doc.add_heading(article.title, level=0)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        f"Rascunho gerado a partir de {cluster.count} tickets recorrentes."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    tickets_para = doc.add_paragraph()
    tickets_para.add_run("Tickets de origem: ").bold = True
    tickets_para.add_run(", ".join(cluster.tickets[:15]))
    if len(cluster.tickets) > 15:
        tickets_para.add_run(f" e mais {len(cluster.tickets) - 15} ticket(s).")

    # ─── Problema ────────────────────────────────────────────────
    doc.add_heading("Problema", level=1)
    doc.add_paragraph(article.problem)

    # ─── Causa raiz ──────────────────────────────────────────────
    doc.add_heading("Causa raiz", level=1)
    doc.add_paragraph(article.cause)

    # ─── Solução ─────────────────────────────────────────────────
    doc.add_heading("Solução", level=1)
    steps = [
        _LEADING_NUMBER_RE.sub("", line).strip()
        for line in article.solution.split("\n")
        if line.strip()
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    # ─── FAQ ─────────────────────────────────────────────────────
    if article.faq:
        doc.add_heading("Perguntas frequentes", level=1)
        for item in article.faq:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(item.question)
            q_run.bold = True
            doc.add_paragraph(item.answer)

    # ─── Metadados ──────────────────────────────────────────────
    doc.add_heading("Metadados", level=1)
    metadata = [
        ("Componentes", ", ".join(cluster.components) or "—"),
        ("Labels", ", ".join(cluster.labels) or "—"),
        ("Tags do artigo", ", ".join(article.tags) or "—"),
        ("Total de tickets do cluster", str(cluster.count)),
    ]
    for label, value in metadata:
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        para.add_run(value)

    # ─── Rodapé com marca ────────────────────────────────────────
    doc.add_paragraph()  # spacer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(SIGNATURE)
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
